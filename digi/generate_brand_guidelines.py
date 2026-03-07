#!/usr/bin/env python3
"""
Digital Web Crew - Brand Guidelines Generator
Comprehensive branding, design system, and motion system
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_color_box(doc, color_hex, color_name, usage):
    """Add a color swatch with description"""
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Color swatch cell
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, color_hex.replace('#', ''))
    cell.width = Inches(0.8)
    
    # Name cell
    cell = table.rows[0].cells[1]
    p = cell.paragraphs[0]
    run = p.add_run(color_name)
    run.font.bold = True
    run.font.size = Pt(11)
    
    # Usage cell
    cell = table.rows[0].cells[2]
    cell.text = usage
    cell.width = Inches(3)
    
    doc.add_paragraph()

def create_doc():
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Inter'
    style.font.size = Pt(11)
    
    # Heading 1 style
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Inter'
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x0a, 0x0a, 0x0f)
    
    # Heading 2 style
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Inter'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    
    # Heading 3 style
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Inter'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x2d, 0x2d, 0x44)
    
    # ===== COVER PAGE =====
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("DIGITAL WEB CREW")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x82, 0xff)
    
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Brand Guidelines")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0a, 0x0a, 0x0f)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = subtitle.add_run("Visual Identity, Design System & Motion Language")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x5a)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = desc.add_run("A technical authority that quietly outperforms everyone else.")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x6a, 0x6a, 0x7a)
    
    doc.add_paragraph()
    
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = version.add_run("Version 1.0 | February 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x8a, 0x8a, 0x9a)
    
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS =====
    doc.add_heading("Contents", level=1)
    
    toc_items = [
        ("1. Brand Strategy & Positioning", 0),
        ("2. Visual Identity System", 0),
        ("   Color System", 1),
        ("   Typography", 1),
        ("   Iconography", 1),
        ("3. UI Design System", 0),
        ("   Spacing & Grid", 1),
        ("   Components", 1),
        ("   States & Feedback", 1),
        ("4. Motion & Animation System", 0),
        ("5. Template System", 0),
        ("6. AI + Design Integration", 0),
        ("7. Implementation Guide", 0),
    ]
    
    for item, level in toc_items:
        p = doc.add_paragraph()
        if level == 1:
            p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(item)
        run.font.size = Pt(11) if level == 0 else Pt(10)
        if level == 0:
            run.font.bold = True
    
    doc.add_page_break()
    
    # ===== SECTION 1: BRAND STRATEGY =====
    doc.add_heading("1. Brand Strategy & Positioning", level=1)
    
    doc.add_heading("Brand Archetype", level=2)
    doc.add_paragraph(
        "PRIMARY: The Ruler - Authority, control, excellence, systems-thinking. "
        "We establish order in chaos and deliver predictable excellence."
    )
    doc.add_paragraph(
        "SECONDARY: The Magician - Transformation, innovation, making the impossible possible. "
        "We turn complex technical challenges into elegant solutions."
    )
    
    doc.add_heading("Brand Personality Traits", level=2)
    traits = [
        ("Precise", "Every detail matters. We measure twice, cut once."),
        ("Confident", "We know our craft and deliver on promises. No fluff."),
        ("Innovative", "AI-first thinking. Always ahead of the curve."),
        ("Reliable", "Systems that work. Every time. Predictable excellence."),
        ("Premium", "Quality over quantity. Craftsmanship in every pixel."),
        ("Transparent", "Clear communication. No hidden complexity."),
    ]
    
    for trait, desc in traits:
        p = doc.add_paragraph()
        run = p.add_run(f"{trait}: ")
        run.font.bold = True
        p.add_run(desc)
    
    doc.add_heading("Emotional Promise", level=2)
    p = doc.add_paragraph()
    run = p.add_run("\"We make the complex effortless.\"")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x00, 0x82, 0xff)
    
    doc.add_paragraph(
        "When clients work with Digital Web Crew, they feel confident that their "
        "digital presence is in expert hands. They experience the relief of "
        "handing off complexity to a team that thrives on it. They see results "
        "that exceed expectations without the drama."
    )
    
    doc.add_heading("What Makes Us Different", level=2)
    doc.add_paragraph("Unlike Vercel partners who focus on implementation alone, we bring:")
    diff = [
        "AI-powered optimization (AEO/GEO) built into every project",
        "System-driven approach with measurable outcomes",
        "Full-stack expertise from strategy to deployment",
        "Performance obsession (Core Web Vitals as baseline)",
        "Productized service model with predictable delivery"
    ]
    for d in diff:
        doc.add_paragraph(d, style='List Bullet')
    
    doc.add_heading("Brand Positioning Statement", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    run = p.add_run(
        "For ambitious companies who demand more than a website, Digital Web Crew "
        "is the AI-powered digital agency that delivers high-performance, "
        "SEO-optimized platforms with measurable business impact. Unlike traditional "
        "agencies, we combine technical excellence with AI-driven optimization to "
        "quietly outperform the competition."
    )
    run.font.italic = True
    
    doc.add_heading("Brand Mantra", level=2)
    p = doc.add_paragraph()
    run = p.add_run("\"Systems over shortcuts. Performance over promises.\"")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x82, 0xff)
    
    doc.add_heading("Do / Don't Rules", level=2)
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    headers = ['DO', 'DON\'T']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    rules = [
        ('Lead with results and metrics', 'Use vague superlatives'),
        ('Show technical depth with clarity', 'Hide behind jargon'),
        ('Use whitespace intentionally', 'Fill space for the sake of it'),
        ('Be confident, not arrogant', 'Be self-deprecating'),
        ('Demonstrate with case studies', 'Make unsubstantiated claims'),
        ('Keep motion purposeful', 'Add animation for decoration'),
    ]
    
    for i, (do, dont) in enumerate(rules, 1):
        table.rows[i].cells[0].text = do
        table.rows[i].cells[1].text = dont
    
    doc.add_page_break()
    
    # ===== SECTION 2: VISUAL IDENTITY =====
    doc.add_heading("2. Visual Identity System", level=1)
    
    doc.add_heading("Color System", level=2)
    
    doc.add_heading("Primary Colors", level=3)
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    
    headers = ['Color', 'Hex', 'RGB', 'Usage']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    colors = [
        ('Deep Navy', '#0A0A0F', '10, 10, 15', 'Primary backgrounds, headers'),
        ('Electric Blue', '#0082FF', '0, 130, 255', 'Primary actions, links, accents'),
        ('Pure White', '#FFFFFF', '255, 255, 255', 'Text on dark, cards'),
    ]
    
    for i, (name, hex_code, rgb, usage) in enumerate(colors, 1):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = hex_code
        table.rows[i].cells[2].text = rgb
        table.rows[i].cells[3].text = usage
    
    doc.add_paragraph()
    
    doc.add_heading("Secondary Colors", level=3)
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    headers = ['Color', 'Hex', 'RGB', 'Usage']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    colors = [
        ('Slate', '#1A1A2E', '26, 26, 46', 'Secondary backgrounds'),
        ('Steel', '#2D2D44', '45, 45, 68', 'Borders, dividers'),
        ('Graphite', '#4A4A5A', '74, 74, 90', 'Secondary text'),
        ('Silver', '#8A8A9A', '138, 138, 154', 'Tertiary text, placeholders'),
    ]
    
    for i, (name, hex_code, rgb, usage) in enumerate(colors, 1):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = hex_code
        table.rows[i].cells[2].text = rgb
        table.rows[i].cells[3].text = usage
    
    doc.add_paragraph()
    
    doc.add_heading("Semantic Colors", level=3)
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    headers = ['Color', 'Hex', 'RGB', 'Usage']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    colors = [
        ('Success', '#00C853', '0, 200, 83', 'Success states, completions'),
        ('Warning', '#FFB300', '255, 179, 0', 'Warnings, attention needed'),
        ('Error', '#FF3D71', '255, 61, 113', 'Errors, critical alerts'),
        ('Info', '#00B0FF', '0, 176, 255', 'Information, neutral alerts'),
    ]
    
    for i, (name, hex_code, rgb, usage) in enumerate(colors, 1):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = hex_code
        table.rows[i].cells[2].text = rgb
        table.rows[i].cells[3].text = usage
    
    doc.add_paragraph()
    
    doc.add_heading("Accessibility Requirements", level=3)
    doc.add_paragraph("All color combinations must meet WCAG AA standards:")
    a11y = [
        "Normal text: 4.5:1 contrast ratio minimum",
        "Large text (18pt+): 3:1 contrast ratio minimum",
        "UI components: 3:1 contrast ratio minimum",
        "Never rely on color alone to convey information"
    ]
    for item in a11y:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Typography", level=2)
    
    doc.add_heading("Font Families", level=3)
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    headers = ['Role', 'Font', 'Usage']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    fonts = [
        ('Headlines', 'Inter (700-900)', 'H1-H3, section titles, CTAs'),
        ('Body/UI', 'Inter (400-600)', 'Body text, UI elements, labels'),
        ('Code/Mono', 'JetBrains Mono', 'Code blocks, technical data'),
    ]
    
    for i, (role, font, usage) in enumerate(fonts, 1):
        table.rows[i].cells[0].text = role
        table.rows[i].cells[1].text = font
        table.rows[i].cells[2].text = usage
    
    doc.add_paragraph()
    
    doc.add_heading("Type Scale", level=3)
    
    table = doc.add_table(rows=9, cols=4)
    table.style = 'Table Grid'
    
    headers = ['Level', 'Size', 'Line Height', 'Weight']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    scale = [
        ('Hero', '64px / 4rem', '1.1', '800'),
        ('H1', '48px / 3rem', '1.2', '700'),
        ('H2', '36px / 2.25rem', '1.25', '700'),
        ('H3', '24px / 1.5rem', '1.3', '600'),
        ('H4', '20px / 1.25rem', '1.4', '600'),
        ('Body Large', '18px / 1.125rem', '1.6', '400'),
        ('Body', '16px / 1rem', '1.6', '400'),
        ('Small', '14px / 0.875rem', '1.5', '400'),
    ]
    
    for i, (level, size, lh, weight) in enumerate(scale, 1):
        table.rows[i].cells[0].text = level
        table.rows[i].cells[1].text = size
        table.rows[i].cells[2].text = lh
        table.rows[i].cells[3].text = weight
    
    doc.add_paragraph()
    
    doc.add_heading("Iconography", level=2)
    doc.add_paragraph(
        "Use Lucide Icons as the primary icon library. Icons should be:"
    )
    icon_rules = [
        "Line style (1.5px stroke weight)",
        "Consistent 24x24px viewbox",
        "Monochrome (inherit text color)",
        "Purposeful - every icon has meaning"
    ]
    for rule in icon_rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== SECTION 3: UI DESIGN SYSTEM =====
    doc.add_heading("3. UI Design System", level=1)
    
    doc.add_heading("Spacing System", level=2)
    doc.add_paragraph("Based on 4px grid system:")
    
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    
    headers = ['Token', 'Value', 'Usage']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    spacing = [
        ('space-1', '4px / 0.25rem', 'Tight gaps'),
        ('space-2', '8px / 0.5rem', 'Icon gaps, small padding'),
        ('space-3', '12px / 0.75rem', 'Button padding'),
        ('space-4', '16px / 1rem', 'Card padding, standard gap'),
        ('space-6', '24px / 1.5rem', 'Section gaps'),
        ('space-8', '32px / 2rem', 'Large section spacing'),
        ('space-12', '48px / 3rem', 'Major section breaks'),
        ('space-16', '64px / 4rem', 'Hero spacing'),
    ]
    
    for i, (token, value, usage) in enumerate(spacing, 1):
        table.rows[i].cells[0].text = token
        table.rows[i].cells[1].text = value
        table.rows[i].cells[2].text = usage
    
    doc.add_paragraph()
    
    doc.add_heading("Grid System", level=2)
    doc.add_paragraph(
        "12-column grid with responsive breakpoints:"
    )
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    
    headers = ['Breakpoint', 'Width', 'Columns/Gutter']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    breakpoints = [
        ('Mobile', '< 640px', '4 columns / 16px gutter'),
        ('Tablet', '640px - 1024px', '8 columns / 24px gutter'),
        ('Desktop', '1024px - 1280px', '12 columns / 24px gutter'),
        ('Large', '1280px - 1536px', '12 columns / 32px gutter'),
        ('XL', '> 1536px', '12 columns / 32px gutter'),
    ]
    
    for i, (bp, width, cols) in enumerate(breakpoints, 1):
        table.rows[i].cells[0].text = bp
        table.rows[i].cells[1].text = width
        table.rows[i].cells[2].text = cols
    
    doc.add_paragraph()
    
    doc.add_heading("Component Styles", level=2)
    
    doc.add_heading("Buttons", level=3)
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    headers = ['Variant', 'Background', 'Text', 'Border']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    buttons = [
        ('Primary', '#0082FF', '#FFFFFF', 'none'),
        ('Secondary', 'transparent', '#0082FF', '1px #0082FF'),
        ('Ghost', 'transparent', '#FFFFFF', 'none'),
        ('Danger', '#FF3D71', '#FFFFFF', 'none'),
    ]
    
    for i, (variant, bg, text, border) in enumerate(buttons, 1):
        table.rows[i].cells[0].text = variant
        table.rows[i].cells[1].text = bg
        table.rows[i].cells[2].text = text
        table.rows[i].cells[3].text = border
    
    doc.add_paragraph()
    doc.add_paragraph("Button sizing:")
    doc.add_paragraph("Small: 32px height, px-3", style='List Bullet')
    doc.add_paragraph("Medium: 40px height, px-4 (default)", style='List Bullet')
    doc.add_paragraph("Large: 48px height, px-6", style='List Bullet')
    
    doc.add_heading("Cards", level=3)
    doc.add_paragraph(
        "Cards use 1px border (#2D2D44) with 8px border-radius. "
        "Background: #0A0A0F or #1A1A2E depending on context. "
        "Padding: 24px (space-6)."
    )
    
    doc.add_heading("Forms", level=3)
    doc.add_paragraph(
        "Input fields: 40px height, 8px radius, 1px border (#2D2D44). "
        "Focus state: border-color #0082FF with 2px glow. "
        "Labels: 14px, font-weight 500, margin-bottom 8px."
    )
    
    doc.add_heading("States & Feedback", level=2)
    
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    
    headers = ['State', 'Visual Treatment', 'Motion']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    states = [
        ('Loading', 'Spinner animation', 'Rotate 360deg, 1s linear'),
        ('Success', 'Green checkmark', 'Scale in 0->1, 200ms'),
        ('Error', 'Red border + message', 'Shake animation, 300ms'),
        ('AI Thinking', 'Pulsing dots', 'Fade in/out sequence'),
        ('Score Ready', 'Number count up', 'Animate from 0, 800ms'),
        ('Hover', 'Subtle lift', 'Translate Y -2px, 150ms'),
    ]
    
    for i, (state, visual, motion) in enumerate(states, 1):
        table.rows[i].cells[0].text = state
        table.rows[i].cells[1].text = visual
        table.rows[i].cells[2].text = motion
    
    doc.add_page_break()
    
    # ===== SECTION 4: MOTION SYSTEM =====
    doc.add_heading("4. Motion & Animation System", level=1)
    
    doc.add_heading("Motion Principles", level=2)
    principles = [
        ("Purposeful", "Every animation serves a function - guiding attention, providing feedback, or indicating state change."),
        ("Fast", "Most animations complete in 150-300ms. Never make users wait."),
        ("Smooth", "Use ease-out for entering elements, ease-in for exiting."),
        ("Restrained", "One animation at a time. No competing motion."),
        ("Respectful", "Honor prefers-reduced-motion. Provide static alternatives."),
    ]
    
    for principle, desc in principles:
        p = doc.add_paragraph()
        run = p.add_run(f"{principle}: ")
        run.font.bold = True
        p.add_run(desc)
    
    doc.add_heading("Easing Functions", level=2)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    headers = ['Name', 'CSS', 'Usage']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    easing = [
        ('Ease Out', 'cubic-bezier(0, 0, 0.2, 1)', 'Elements entering'),
        ('Ease In', 'cubic-bezier(0.4, 0, 1, 1)', 'Elements exiting'),
        ('Ease In Out', 'cubic-bezier(0.4, 0, 0.2, 1)', 'State changes'),
        ('Spring', 'cubic-bezier(0.34, 1.56, 0.64, 1)', 'Playful interactions'),
    ]
    
    for i, (name, css, usage) in enumerate(easing, 1):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = css
        table.rows[i].cells[2].text = usage
    
    doc.add_paragraph()
    
    doc.add_heading("Animation Patterns", level=2)
    
    doc.add_heading("Page Transitions", level=3)
    doc.add_paragraph(
        "Use fade + slight upward movement (8px) for page transitions. "
        "Duration: 200ms, easing: ease-out."
    )
    
    doc.add_heading("Scroll Animations", level=3)
    doc.add_paragraph(
        "Elements fade in and slide up (16px) as they enter viewport. "
        "Trigger at 20% visibility. Stagger multiple elements by 50ms."
    )
    
    doc.add_heading("Micro-interactions", level=3)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    headers = ['Element', 'Trigger', 'Animation']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    micro = [
        ('Button', 'Hover', 'Scale 1.02, shadow increase'),
        ('Button', 'Click', 'Scale 0.98, quick release'),
        ('Card', 'Hover', 'Translate Y -4px, shadow increase'),
        ('Input', 'Focus', 'Border color change, subtle glow'),
    ]
    
    for i, (elem, trigger, anim) in enumerate(micro, 1):
        table.rows[i].cells[0].text = elem
        table.rows[i].cells[1].text = trigger
        table.rows[i].cells[2].text = anim
    
    doc.add_paragraph()
    
    doc.add_heading("AI-Specific Animations", level=3)
    doc.add_paragraph(
        "AI thinking state: Three pulsing dots with staggered fade (0.6s cycle). "
        "Content generation: Progress bar with shimmer effect. "
        "Score reveal: Number count-up animation (800ms)."
    )
    
    doc.add_heading("Performance Constraints", level=2)
    perf = [
        "Use transform and opacity only (GPU accelerated)",
        "Avoid animating layout properties (width, height, top, left)",
        "Use will-change sparingly and remove after animation",
        "Respect prefers-reduced-motion media query",
        "Test on low-end devices"
    ]
    for item in perf:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== SECTION 5: TEMPLATES =====
    doc.add_heading("5. Template System", level=1)
    
    templates = [
        {
            'name': 'Homepage',
            'sections': [
                'Hero: Value prop + CTA + social proof',
                'Services: 3-4 key offerings with icons',
                'Process: How we work (3-4 steps)',
                'Case Studies: 2-3 featured projects',
                'Testimonials: Client quotes',
                'CTA: Final conversion point'
            ],
            'seo_blocks': ['Schema.org Organization', 'FAQ schema', 'Service schema']
        },
        {
            'name': 'Service Page',
            'sections': [
                'Hero: Service name + benefit statement',
                'Problem: Pain points we solve',
                'Solution: Our approach/methodology',
                'Features: Detailed capabilities',
                'Case Study: Relevant example',
                'Pricing/CTA: Next steps'
            ],
            'seo_blocks': ['Service schema', 'BreadcrumbList', 'HowTo schema']
        },
        {
            'name': 'Case Study',
            'sections': [
                'Hero: Client name + results',
                'Challenge: What needed solving',
                'Approach: Our methodology',
                'Solution: What we built',
                'Results: Metrics and impact',
                'Testimonial: Client quote'
            ],
            'seo_blocks': ['Article schema', 'Review schema']
        },
        {
            'name': 'Blog Post (AEO Optimized)',
            'sections': [
                'H1: Question or direct answer',
                'Quick Answer: 40-60 word summary',
                'Table of Contents: Jump links',
                'Body: Detailed explanation',
                'FAQ: Related questions',
                'CTA: Related service/content'
            ],
            'seo_blocks': ['Article schema', 'FAQPage schema', 'BreadcrumbList']
        },
        {
            'name': 'GEO Landing Page',
            'sections': [
                'Hero: Location-specific headline',
                'Local Context: Why this location matters',
                'Services: Location-relevant offerings',
                'Local Proof: Regional clients/case studies',
                'Team: Local presence/expertise',
                'Contact: Location-specific CTA'
            ],
            'seo_blocks': ['LocalBusiness schema', 'Service schema']
        },
        {
            'name': 'Admin Dashboard',
            'sections': [
                'Sidebar: Navigation menu',
                'Header: Page title + actions',
                'Stats Row: Key metrics cards',
                'Main Content: Tables/charts/forms',
                'Filters: Search and filter controls'
            ],
            'seo_blocks': []
        },
    ]
    
    for template in templates:
        doc.add_heading(template['name'], level=2)
        doc.add_paragraph("Sections:")
        for section in template['sections']:
            doc.add_paragraph(section, style='List Bullet')
        if template['seo_blocks']:
            doc.add_paragraph("SEO/AEO Blocks:")
            for block in template['seo_blocks']:
                doc.add_paragraph(block, style='List Bullet')
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ===== SECTION 6: AI + DESIGN =====
    doc.add_heading("6. AI + Design Integration", level=1)
    
    doc.add_heading("AI Content in Design", level=2)
    doc.add_paragraph(
        "AI-generated content must fit seamlessly into our design system:"
    )
    ai_content = [
        "Content length variations handled by flexible layouts",
        "Dynamic content areas use min/max height constraints",
        "Overflow handled gracefully with fade or scroll",
        "Loading states for AI-generated content",
        "Error states when AI fails to generate"
    ]
    for item in ai_content:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("AI Score Visualization", level=2)
    doc.add_paragraph(
        "When displaying AI scores (SEO, readability, etc.):"
    )
    score_viz = [
        "Use circular progress indicators for single scores",
        "Use horizontal bars for comparative metrics",
        "Color-code: Green (80-100), Yellow (60-79), Red (<60)",
        "Animate score reveal with count-up effect",
        "Show improvement suggestions alongside scores"
    ]
    for item in score_viz:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("AI Suggestion UI", level=2)
    doc.add_paragraph(
        "AI suggestions should be actionable and clear:"
    )
    suggestions = [
        "Show suggestion type icon (content, SEO, grammar)",
        "Display before/after comparison when applicable",
        "One-click apply for safe changes",
        "Review required for significant changes",
        "Batch apply option for multiple suggestions"
    ]
    for item in suggestions:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Design Guardrails", level=2)
    doc.add_paragraph(
        "Prevent AI from breaking design consistency:"
    )
    guardrails = [
        "AI cannot modify brand colors or typography",
        "AI-generated images must match brand style",
        "Content must fit within established layouts",
        "All AI output reviewed before publication",
        "Fallback designs for AI failure states"
    ]
    for item in guardrails:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== SECTION 7: IMPLEMENTATION =====
    doc.add_heading("7. Implementation Guide", level=1)
    
    doc.add_heading("Tailwind Configuration", level=2)
    doc.add_paragraph("Extend your tailwind.config.ts:")
    
    config_code = '''
// tailwind.config.ts
import type { Config } from "tailwindcss";

const config: Config = {
  theme: {
    extend: {
      colors: {
        // Primary
        navy: {
          DEFAULT: '#0A0A0F',
          light: '#1A1A2E',
        },
        blue: {
          DEFAULT: '#0082FF',
          dark: '#0066CC',
          light: '#4DA3FF',
        },
        // Secondary
        slate: '#1A1A2E',
        steel: '#2D2D44',
        graphite: '#4A4A5A',
        silver: '#8A8A9A',
        // Semantic
        success: '#00C853',
        warning: '#FFB300',
        error: '#FF3D71',
        info: '#00B0FF',
      },
      fontFamily: {
        sans: ['Inter', 'system-ui', 'sans-serif'],
        mono: ['JetBrains Mono', 'monospace'],
      },
      spacing: {
        '18': '4.5rem',
        '22': '5.5rem',
      },
      animation: {
        'fade-in': 'fadeIn 200ms ease-out',
        'slide-up': 'slideUp 200ms ease-out',
        'pulse-slow': 'pulse 3s infinite',
        'spin-slow': 'spin 3s linear infinite',
      },
      keyframes: {
        fadeIn: {
          '0%': { opacity: '0' },
          '100%': { opacity: '1' },
        },
        slideUp: {
          '0%': { opacity: '0', transform: 'translateY(8px)' },
          '100%': { opacity: '1', transform: 'translateY(0)' },
        },
      },
    },
  },
};
'''
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(config_code)
    run.font.name = 'JetBrains Mono'
    run.font.size = Pt(9)
    
    doc.add_heading("CSS Variables", level=2)
    doc.add_paragraph("Define in globals.css:")
    
    css_code = '''
:root {
  /* Colors */
  --color-navy: #0A0A0F;
  --color-navy-light: #1A1A2E;
  --color-blue: #0082FF;
  --color-blue-dark: #0066CC;
  --color-steel: #2D2D44;
  --color-graphite: #4A4A5A;
  --color-silver: #8A8A9A;
  
  /* Spacing */
  --space-1: 0.25rem;
  --space-2: 0.5rem;
  --space-3: 0.75rem;
  --space-4: 1rem;
  --space-6: 1.5rem;
  --space-8: 2rem;
  
  /* Animation */
  --duration-fast: 150ms;
  --duration-normal: 200ms;
  --duration-slow: 300ms;
  --ease-out: cubic-bezier(0, 0, 0.2, 1);
  --ease-in: cubic-bezier(0.4, 0, 1, 1);
  --ease-in-out: cubic-bezier(0.4, 0, 0.2, 1);
}
'''
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(css_code)
    run.font.name = 'JetBrains Mono'
    run.font.size = Pt(9)
    
    doc.add_heading("Component Examples", level=2)
    doc.add_paragraph("Primary Button:")
    
    btn_code = '''
// components/ui/Button.tsx
import { cn } from "@/lib/utils";

interface ButtonProps {
  variant?: 'primary' | 'secondary' | 'ghost' | 'danger';
  size?: 'sm' | 'md' | 'lg';
  children: React.ReactNode;
  className?: string;
}

export function Button({
  variant = 'primary',
  size = 'md',
  children,
  className,
  ...props
}: ButtonProps) {
  return (
    <button
      className={cn(
        // Base
        "inline-flex items-center justify-center font-medium rounded-lg",
        "transition-all duration-200 ease-out",
        "focus:outline-none focus:ring-2 focus:ring-blue/50",
        "disabled:opacity-50 disabled:cursor-not-allowed",
        
        // Variants
        variant === 'primary' && [
          "bg-blue text-white",
          "hover:bg-blue-dark hover:scale-[1.02]",
          "active:scale-[0.98]",
        ],
        variant === 'secondary' && [
          "bg-transparent text-blue border border-blue",
          "hover:bg-blue/10",
        ],
        variant === 'ghost' && [
          "bg-transparent text-white",
          "hover:bg-white/10",
        ],
        variant === 'danger' && [
          "bg-error text-white",
          "hover:opacity-90",
        ],
        
        // Sizes
        size === 'sm' && "h-8 px-3 text-sm",
        size === 'md' && "h-10 px-4 text-base",
        size === 'lg' && "h-12 px-6 text-lg",
        
        className
      )}
      {...props}
    >
      {children}
    </button>
  );
}
'''
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(btn_code)
    run.font.name = 'JetBrains Mono'
    run.font.size = Pt(9)
    
    doc.add_heading("Design Tokens JSON", level=2)
    doc.add_paragraph("Export for design tools:")
    
    tokens_code = '''
{
  "colors": {
    "navy": { "DEFAULT": "#0A0A0F", "light": "#1A1A2E" },
    "blue": { "DEFAULT": "#0082FF", "dark": "#0066CC", "light": "#4DA3FF" },
    "steel": "#2D2D44",
    "graphite": "#4A4A5A",
    "silver": "#8A8A9A",
    "success": "#00C853",
    "warning": "#FFB300",
    "error": "#FF3D71",
    "info": "#00B0FF"
  },
  "typography": {
    "fontFamily": {
      "sans": ["Inter", "system-ui", "sans-serif"],
      "mono": ["JetBrains Mono", "monospace"]
    },
    "sizes": {
      "hero": { "size": "4rem", "lineHeight": "1.1", "weight": "800" },
      "h1": { "size": "3rem", "lineHeight": "1.2", "weight": "700" },
      "h2": { "size": "2.25rem", "lineHeight": "1.25", "weight": "700" },
      "h3": { "size": "1.5rem", "lineHeight": "1.3", "weight": "600" },
      "body": { "size": "1rem", "lineHeight": "1.6", "weight": "400" }
    }
  },
  "spacing": {
    "1": "0.25rem",
    "2": "0.5rem",
    "3": "0.75rem",
    "4": "1rem",
    "6": "1.5rem",
    "8": "2rem",
    "12": "3rem",
    "16": "4rem"
  },
  "animation": {
    "duration": {
      "fast": "150ms",
      "normal": "200ms",
      "slow": "300ms"
    },
    "easing": {
      "out": "cubic-bezier(0, 0, 0.2, 1)",
      "in": "cubic-bezier(0.4, 0, 1, 1)",
      "inOut": "cubic-bezier(0.4, 0, 0.2, 1)"
    }
  }
}
'''
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(tokens_code)
    run.font.name = 'JetBrains Mono'
    run.font.size = Pt(9)
    
    # Save document
    output_path = "/mnt/okcomputer/output/DigitalWebCrew_Brand_Guidelines.docx"
    doc.save(output_path)
    print(f"Brand Guidelines saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_doc()
