#!/usr/bin/env python3
"""
Generate Technical Specification Document for DigiWebCrew.com
Using python-docx for reliable document generation
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_code_block(doc, code_text):
    """Add a styled code block"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    
    # Add shading via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'f5f7fa')
    pPr.append(shd)

def create_doc():
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Heading 1 style
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
    
    # Heading 2 style
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2d, 0x5a, 0x87)
    
    # Heading 3 style
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x3d, 0x7a, 0xb8)
    
    # ===== COVER PAGE =====
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Technical Specification")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0f, 0x1f, 0x33)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = subtitle.add_run("DigiWebCrew.com Platform")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
    
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = desc.add_run("Full-Stack Architecture Blueprint, Implementation Plan & Developer Guide")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x5a)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = company.add_run("Digital Web Crew")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
    
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = version.add_run("Version 1.0 | February 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x7a, 0x7a, 0x8a)
    
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = conf.add_run("Confidential - For Internal Development Team")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x7a, 0x7a, 0x8a)
    
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS =====
    toc_title = doc.add_heading("Table of Contents", level=1)
    
    toc_items = [
        ("Executive Summary", 0),
        ("Technology Stack & Architecture", 0),
        ("  Frontend Architecture", 1),
        ("  Backend & API Layer", 1),
        ("  Database Design", 1),
        ("SEO / AEO / GEO Strategy", 0),
        ("  On-Page SEO Framework", 1),
        ("  AEO Implementation", 1),
        ("  GEO Targeting System", 1),
        ("AI Content & Publishing Framework", 0),
        ("Chatbot & AI Agents Implementation", 0),
        ("Admin Panels & Internal Tools", 0),
        ("Marketing Automation & Lead Systems", 0),
        ("CI/CD, Testing & Deployment", 0),
        ("Security Architecture", 0),
        ("Developer Task Breakdown", 0),
        ("Appendices", 0),
    ]
    
    for item, level in toc_items:
        p = doc.add_paragraph()
        if level == 1:
            p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(item)
        run.font.size = Pt(11) if level == 0 else Pt(10)
        if level == 0:
            run.font.bold = True
    
    doc.add_page_break()
    
    # ===== EXECUTIVE SUMMARY =====
    doc.add_heading("Executive Summary", level=1)
    
    doc.add_paragraph(
        "This Technical Specification Document provides a comprehensive blueprint for building, "
        "deploying, and maintaining the DigiWebCrew.com digital agency platform. The document "
        "serves as the single source of truth for developers, architects, and stakeholders involved "
        "in the project."
    )
    
    doc.add_heading("Project Overview", level=2)
    doc.add_paragraph(
        "DigiWebCrew.com is a next-generation digital agency platform designed to showcase services, "
        "capture leads, and demonstrate technical excellence. The platform incorporates cutting-edge "
        "technologies including AI-powered content generation, intelligent chatbots, and advanced SEO "
        "optimization strategies."
    )
    
    doc.add_heading("Key Objectives", level=2)
    objectives = [
        "Build a high-performance, SEO-optimized web platform using Next.js 15+ with App Router",
        "Implement AI-powered content generation and publishing workflows",
        "Deploy intelligent chatbot systems for lead capture and customer engagement",
        "Establish comprehensive SEO/AEO/GEO optimization frameworks",
        "Create robust admin panels for content and business management",
        "Implement marketing automation and lead nurturing systems",
        "Ensure enterprise-grade security, testing, and deployment practices"
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading("Target Audience", level=2)
    doc.add_paragraph("This document is intended for:")
    audiences = [
        "Full-stack developers implementing features and integrations",
        "DevOps engineers managing deployment and infrastructure",
        "SEO specialists optimizing content and technical performance",
        "Project managers tracking progress and deliverables",
        "Technical architects reviewing system design decisions"
    ]
    for aud in audiences:
        doc.add_paragraph(aud, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== TECHNOLOGY STACK =====
    doc.add_heading("Technology Stack & Architecture", level=1)
    
    doc.add_paragraph(
        "The DigiWebCrew platform is built on a modern, scalable technology stack designed for "
        "performance, maintainability, and developer experience."
    )
    
    doc.add_heading("Frontend Architecture", level=2)
    doc.add_heading("Core Technologies", level=3)
    
    # Tech table
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    
    headers = ['Technology', 'Version/Tool', 'Purpose']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    tech_data = [
        ['Framework', 'Next.js 16.1.1', 'React framework with SSR/SSG/ISR'],
        ['Language', 'TypeScript 5.x', 'Type-safe JavaScript'],
        ['Styling', 'Tailwind CSS 4.x', 'Utility-first CSS framework'],
        ['Animation', 'Framer Motion 12.x', 'React animation library'],
        ['UI Components', 'Radix UI + shadcn/ui', 'Accessible components'],
        ['Icons', 'Lucide React', 'Modern icon library'],
        ['State', 'React Hooks + Context', 'Built-in state management']
    ]
    
    for i, row_data in enumerate(tech_data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading("Backend & API Layer", level=2)
    doc.add_paragraph(
        "The backend leverages Next.js Route Handlers for API endpoints, providing a unified "
        "full-stack development experience."
    )
    
    backend_features = [
        "RESTful API design with consistent response formats",
        "Middleware-based authentication and authorization",
        "Rate limiting and request validation with Zod",
        "Error handling with structured error responses",
        "API versioning strategy for backward compatibility"
    ]
    for feat in backend_features:
        doc.add_paragraph(feat, style='List Bullet')
    
    doc.add_heading("Database Design", level=2)
    doc.add_paragraph(
        "The platform uses MongoDB with Mongoose ODM for flexible document-based data storage."
    )
    
    db_benefits = [
        "Schema flexibility for evolving content models",
        "Native JSON support for modern web applications",
        "Horizontal scaling capabilities for growth",
        "Rich query capabilities with aggregation pipelines"
    ]
    for benefit in db_benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== SEO STRATEGY =====
    doc.add_heading("SEO / AEO / GEO Strategy", level=1)
    
    doc.add_paragraph(
        "Search Engine Optimization (SEO), Answer Engine Optimization (AEO), and Generative Engine "
        "Optimization (GEO) form the foundation of the platform's visibility strategy."
    )
    
    doc.add_heading("On-Page SEO Framework", level=2)
    doc.add_heading("Dynamic Metadata Generation", level=3)
    doc.add_paragraph(
        "Next.js 15 provides powerful metadata APIs for SEO optimization including dynamic titles, "
        "descriptions, Open Graph tags, and structured data implementation."
    )
    
    doc.add_heading("Structured Data Implementation", level=3)
    doc.add_paragraph(
        "JSON-LD structured data is implemented for rich search results including Organization schema, "
        "Service schema, LocalBusiness schema, and Article schema for blog posts."
    )
    
    doc.add_heading("AEO Implementation", level=2)
    doc.add_paragraph(
        "Answer Engine Optimization requires content structured for AI comprehension with focus on:"
    )
    
    aeo_factors = [
        "Information Gain - Unique insights and original research",
        "Entity Coverage - Comprehensive topic entities",
        "Question Coverage - FAQ sections and Q&A format",
        "Citation Worthiness - Authoritative sources and data",
        "Readability - Clear structure and scannable content"
    ]
    for factor in aeo_factors:
        doc.add_paragraph(factor, style='List Bullet')
    
    doc.add_heading("GEO Targeting System", level=2)
    doc.add_paragraph(
        "Generative Engine Optimization and geo-targeting work together for location-based visibility "
        "with region-specific content variations and LocalBusiness schema implementation."
    )
    
    doc.add_page_break()
    
    # ===== AI CONTENT FRAMEWORK =====
    doc.add_heading("AI Content & Publishing Framework", level=1)
    
    doc.add_paragraph(
        "The AI Content Framework enables automated content generation, optimization, and publishing "
        "while maintaining quality control and brand consistency."
    )
    
    doc.add_heading("Content Generation Workflows", level=2)
    doc.add_paragraph("The content workflow follows a structured pipeline:")
    
    workflow_steps = [
        "Topic Generation - AI analyzes trends to suggest high-value topics",
        "Outline Creation - AI generates comprehensive content outlines",
        "Draft Generation - AI produces full draft content",
        "Human Review - Editor reviews and approves content",
        "SEO Optimization - AI enhances metadata and keyword density",
        "Publication - Content is scheduled and published"
    ]
    for step in workflow_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading("Version Control & Rollback", level=2)
    doc.add_paragraph(
        "Content version history enables safe experimentation with automatic versioning on every save, "
        "diff comparison between versions, and one-click rollback functionality."
    )
    
    doc.add_page_break()
    
    # ===== CHATBOT =====
    doc.add_heading("Chatbot & AI Agents Implementation", level=1)
    
    doc.add_paragraph(
        "The intelligent chatbot system serves as a 24/7 virtual assistant for lead capture, "
        "customer support, and guided sales conversations."
    )
    
    doc.add_heading("Functional Requirements", level=2)
    
    # Chatbot features table
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    
    headers = ['Capability', 'Description', 'Priority']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    chatbot_data = [
        ['Lead Capture', 'Collect visitor information', 'High'],
        ['Quote Assistance', 'Guide service selection', 'High'],
        ['FAQ Handling', 'Answer common questions', 'High'],
        ['Recommendations', 'Suggest relevant services', 'Medium'],
        ['Appointment Booking', 'Schedule consultations', 'Medium'],
        ['CRM Integration', 'Sync with HubSpot', 'High']
    ]
    
    for i, row_data in enumerate(chatbot_data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading("Hybrid AI Architecture", level=2)
    doc.add_paragraph(
        "The chatbot uses a hybrid approach combining rule-based responses for common queries "
        "and AI-powered responses for complex conversations."
    )
    
    doc.add_heading("Data & Security", level=2)
    security_items = [
        "All conversations encrypted in transit (TLS 1.3) and at rest (AES-256)",
        "PII automatically detected and masked in logs",
        "Conversation data retained for 90 days then anonymized",
        "GDPR-compliant data handling with right-to-erasure support"
    ]
    for item in security_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== ADMIN PANELS =====
    doc.add_heading("Admin Panels & Internal Tools", level=1)
    
    doc.add_paragraph(
        "The admin dashboard provides comprehensive tools for content management, SEO optimization, "
        "AI workflows, and business analytics."
    )
    
    doc.add_heading("Content Management", level=2)
    doc.add_paragraph("Features include:")
    content_features = [
        "Content Calendar with monthly/weekly views",
        "Drag-and-drop scheduling",
        "Status indicators for content state",
        "Team assignments and workflow management"
    ]
    for feat in content_features:
        doc.add_paragraph(feat, style='List Bullet')
    
    doc.add_heading("SEO Optimization Panel", level=2)
    doc.add_paragraph(
        "Real-time SEO scoring with suggestions for title optimization, meta description improvements, "
        "keyword density analysis, and readability scoring."
    )
    
    doc.add_heading("Analytics Dashboard", level=2)
    doc.add_paragraph(
        "Key metrics visualization including traffic analytics, lead conversion rates, chatbot "
        "interaction metrics, and content performance tracking."
    )
    
    doc.add_page_break()
    
    # ===== MARKETING AUTOMATION =====
    doc.add_heading("Marketing Automation & Lead Systems", level=1)
    
    doc.add_paragraph(
        "The marketing automation system streamlines lead capture, nurturing, and conversion through "
        "integrated workflows and intelligent triggers."
    )
    
    doc.add_heading("CRM Integration", level=2)
    doc.add_paragraph(
        "HubSpot integration enables automatic lead synchronization, contact enrichment, "
        "note creation, and deal pipeline management."
    )
    
    doc.add_heading("Email Automation", level=2)
    doc.add_paragraph(
        "Drip campaign workflows nurture leads through the sales funnel with welcome sequences, "
        "educational content, and consultation offers."
    )
    
    doc.add_heading("Tagging & Segmentation", level=2)
    doc.add_paragraph("Dynamic segments include:")
    segments = [
        "Hot Leads - Visited pricing page + spent >3 minutes",
        "Enterprise - Company size >500 employees",
        "SMB - Company size <100 employees",
        "Content Engaged - Downloaded 2+ resources",
        "Chatbot Qualified - Completed lead capture flow"
    ]
    for seg in segments:
        doc.add_paragraph(seg, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== CI/CD =====
    doc.add_heading("CI/CD, Testing & Deployment", level=1)
    
    doc.add_paragraph(
        "The continuous integration and deployment pipeline ensures code quality, automated testing, "
        "and reliable production releases."
    )
    
    doc.add_heading("CI/CD Pipeline", level=2)
    doc.add_paragraph("GitHub Actions workflow includes:")
    cicd_steps = [
        "Lint and TypeScript type checking",
        "Unit and integration test execution",
        "Build verification",
        "Automatic deployment to staging (develop branch)",
        "Production deployment with smoke tests (main branch)"
    ]
    for step in cicd_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_heading("Testing Strategy", level=2)
    
    # Testing table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    
    headers = ['Test Type', 'Framework', 'Coverage', 'Command']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    test_data = [
        ['Unit Tests', 'Jest', '80%', 'npm run test:unit'],
        ['Integration', 'Jest + MSW', '60%', 'npm run test:integration'],
        ['E2E Tests', 'Playwright', 'Critical paths', 'npm run test:e2e'],
        ['Visual', 'Chromatic', 'UI components', 'npm run test:visual'],
        ['Performance', 'Lighthouse CI', 'CWV', 'npm run test:perf']
    ]
    
    for i, row_data in enumerate(test_data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading("Production Safety", level=2)
    doc.add_paragraph("Pre-deployment checklist:")
    checklist = [
        "All tests passing (unit, integration, E2E)",
        "Lighthouse scores: Performance >90, Accessibility >95",
        "No security vulnerabilities (npm audit)",
        "Environment variables configured",
        "Database migrations applied"
    ]
    for item in checklist:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== SECURITY =====
    doc.add_heading("Security Architecture", level=1)
    
    doc.add_paragraph(
        "Security is integrated into every layer of the application following industry best practices."
    )
    
    # Security table
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    
    headers = ['Layer', 'Measures', 'Tools']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    security_data = [
        ['Application', 'Input validation, XSS protection', 'Zod, Helmet'],
        ['Authentication', 'JWT with refresh tokens', 'NextAuth.js'],
        ['Authorization', 'RBAC, permissions', 'CASL'],
        ['Data', 'Encryption at rest/transit', 'AES-256, TLS 1.3'],
        ['API', 'Rate limiting, validation', 'Express Rate Limit'],
        ['Infrastructure', 'WAF, DDoS protection', 'Cloudflare']
    ]
    
    for i, row_data in enumerate(security_data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading("Dependency Management", level=2)
    doc.add_paragraph(
        "Regular dependency updates with npm audit, automated Dependabot alerts, "
        "Snyk integration for vulnerability scanning, and lockfile maintenance."
    )
    
    doc.add_page_break()
    
    # ===== TASK BREAKDOWN =====
    doc.add_heading("Developer Task Breakdown", level=1)
    
    doc.add_paragraph(
        "Detailed breakdown of development tasks with estimated hours for project planning."
    )
    
    phases = [
        ("Phase 1: Foundation", ["Setup Repository", "Configure TypeScript/ESLint", "Setup Tailwind", 
                                  "Configure shadcn/ui", "Environment Config", "MongoDB Setup", 
                                  "NextAuth.js Config", "Base Layout"], 31),
        ("Phase 2: Core Features", ["Homepage", "Services Page", "About Page", "Portfolio", 
                                    "Contact Form", "Blog Listing", "Blog Detail", "Navigation"], 72),
        ("Phase 3: SEO & Content", ["Metadata System", "Sitemap", "Structured Data", "Canonical URLs",
                                     "Open Graph", "AEO Scoring", "Internal Links", "GEO Pages"], 53),
        ("Phase 4: AI Integration", ["OpenAI Integration", "Content Workflow", "Review Queue",
                                      "Prompt Templates", "Version Control", "SEO AI", "Image Gen"], 60),
        ("Phase 5: Chatbot", ["Chatbot UI", "Intent Detection", "Conversation Flows", "Lead Capture",
                               "AI Responses", "Chat History", "CRM Integration"], 56),
        ("Phase 6: Admin Panel", ["Admin Layout", "Dashboard", "Content Management", "Calendar",
                                   "SEO Panel", "AI Queue", "Chatbot Training", "Analytics", 
                                   "Lead Management", "User Management"], 94),
        ("Phase 7: Marketing", ["HubSpot Integration", "Email Setup", "Templates", "Drip Campaigns",
                                 "Segmentation", "Event Tracking", "Workflows"], 62),
        ("Phase 8: Testing & Deploy", ["Unit Tests", "Integration Tests", "E2E Tests", "CI/CD",
                                         "Vercel Config", "Performance", "Security Audit", "Docs"], 70)
    ]
    
    for phase_name, tasks, hours in phases:
        doc.add_heading(phase_name, level=2)
        for task in tasks:
            doc.add_paragraph(task, style='List Bullet')
        p = doc.add_paragraph()
        p.add_run(f"Estimated Hours: {hours}").bold = True
        doc.add_paragraph()
    
    # Total
    total = sum(h for _, _, h in phases)
    p = doc.add_paragraph()
    run = p.add_run(f"TOTAL ESTIMATED HOURS: {total} (~12 weeks)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
    
    doc.add_page_break()
    
    # ===== APPENDICES =====
    doc.add_heading("Appendices", level=1)
    
    doc.add_heading("Appendix A: API Endpoints", level=2)
    
    # API table
    table = doc.add_table(rows=14, cols=4)
    table.style = 'Table Grid'
    
    headers = ['Endpoint', 'Method', 'Description', 'Auth']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    api_data = [
        ['/api/auth/[...nextauth]', 'ALL', 'Authentication', 'Public'],
        ['/api/content', 'GET', 'List content', 'Optional'],
        ['/api/content', 'POST', 'Create content', 'Admin'],
        ['/api/content/[id]', 'GET', 'Get content', 'Optional'],
        ['/api/content/[id]', 'PUT', 'Update content', 'Admin'],
        ['/api/content/[id]', 'DELETE', 'Delete content', 'Admin'],
        ['/api/leads', 'GET', 'List leads', 'Admin'],
        ['/api/leads', 'POST', 'Create lead', 'Public'],
        ['/api/chat', 'POST', 'Chatbot message', 'Public'],
        ['/api/chat/session', 'POST', 'Create session', 'Public'],
        ['/api/ai/generate', 'POST', 'Generate AI content', 'Admin'],
        ['/api/seo/analyze', 'POST', 'Analyze SEO', 'Admin'],
        ['/api/analytics', 'GET', 'Get analytics', 'Admin']
    ]
    
    for i, row_data in enumerate(api_data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading("Appendix B: Environment Variables", level=2)
    env_vars = [
        "MONGODB_URI - MongoDB connection string",
        "NEXTAUTH_URL - Application URL",
        "NEXTAUTH_SECRET - Auth secret (min 32 chars)",
        "GOOGLE_CLIENT_ID - OAuth client ID",
        "GOOGLE_CLIENT_SECRET - OAuth client secret",
        "OPENAI_API_KEY - OpenAI API key",
        "ANTHROPIC_API_KEY - Anthropic API key",
        "HUBSPOT_ACCESS_TOKEN - HubSpot API token",
        "GOOGLE_ANALYTICS_ID - GA tracking ID",
        "ENCRYPTION_KEY - Data encryption key"
    ]
    for var in env_vars:
        doc.add_paragraph(var, style='List Bullet')
    
    doc.add_heading("Appendix C: Glossary", level=2)
    
    # Glossary table
    table = doc.add_table(rows=9, cols=2)
    table.style = 'Table Grid'
    
    headers = ['Term', 'Definition']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'e8eef4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    glossary_data = [
        ['AEO', 'Answer Engine Optimization - optimizing for AI search'],
        ['GEO', 'Generative Engine Optimization - optimizing for generative AI'],
        ['ISR', 'Incremental Static Regeneration - Next.js feature'],
        ['RAG', 'Retrieval-Augmented Generation - AI technique'],
        ['RBAC', 'Role-Based Access Control - permission system'],
        ['SSG', 'Static Site Generation - pre-rendering at build'],
        ['SSR', 'Server-Side Rendering - rendering on request'],
        ['CWV', 'Core Web Vitals - Google performance metrics']
    ]
    
    for i, row_data in enumerate(glossary_data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    # Save document
    output_path = "/mnt/okcomputer/output/DigiWebCrew_Technical_Specification.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_doc()
